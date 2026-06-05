import os, csv, json, zipfile, shutil
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

BASE = os.environ.get('ROBO_PROJECT_ROOT', os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
OUT=os.path.join(BASE,'outputs')
DOCX=os.path.join(BASE,'manuscript','robo_advisor_final_submission_ready.docx')

# ---------------------- helpers ----------------------
def read_csv(name):
    path=os.path.join(OUT,name)
    with open(path,encoding='utf-8-sig',newline='') as f:
        return list(csv.DictReader(f))

def fmt_num(x, dec=2):
    try:
        x=float(x)
        if abs(x) < 0.001 and x != 0:
            return f'{x:.2e}'
        return f'{x:.{dec}f}'
    except Exception:
        return str(x)

def pstar(p):
    try:
        p=float(p)
        if p < 0.001: return '***'
        if p < 0.01: return '**'
        if p < 0.05: return '*'
        return ''
    except Exception:
        return ''

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def shade_cell(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)

def set_cell_text(cell, text, bold=False, size=9, color=None):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(str(text))
    r.bold = bold
    r.font.size = Pt(size)
    r.font.name = 'Times New Roman'
    if color:
        r.font.color.rgb = RGBColor(*color)

def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(10)
    p.paragraph_format.space_after = Pt(8)
    return p

def add_fig(doc, img, caption, width=6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(os.path.join(OUT,img), width=Inches(width))
    add_caption(doc, caption)

def add_table(doc, rows, headers, title, note=None, widths=None, font_size=8):
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(11)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i,h in enumerate(headers):
        set_cell_text(hdr[i], h, bold=True, size=font_size, color=(255,255,255))
        shade_cell(hdr[i], '2D4059')
    for row in rows:
        cells = table.add_row().cells
        for i,h in enumerate(headers):
            set_cell_text(cells[i], row.get(h,''), size=font_size)
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in table.rows:
            for idx,w in enumerate(widths):
                row.cells[idx].width = Inches(w)
    if note:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        rr = p.add_run('Note. ' + note)
        rr.italic = True
        rr.font.size = Pt(9)
    return table

def add_body_p(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(12)
    r.font.name = 'Times New Roman'
    return p

def add_heading(doc, text, level=1):
    p=doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.name='Times New Roman'
        r.font.color.rgb = RGBColor(45,64,89) if level==1 else RGBColor(0,119,182)
    return p

# ---------------------- load outputs ----------------------
with open(os.path.join(OUT,'analysis_summary.json'),encoding='utf-8') as f:
    summary=json.load(f)
rel=read_csv('study2_construct_reliability.csv')
pls=read_csv('study2_pls_bootstrap_paths.csv')
ols_wu=read_csv('study2_ols_wu.csv')
ols_risk=read_csv('study2_ols_risk.csv')
ml_perf=read_csv('combined_ml_performance.csv')
nfcs_ind=read_csv('study3_nfcs_weighted_indicators.csv')
nfcs_wls=read_csv('study3_nfcs_wls_digital_trading.csv')

# ---------------------- doc setup ----------------------
doc=Document()
sec=doc.sections[0]
sec.top_margin = Inches(1.0)
sec.bottom_margin = Inches(1.0)
sec.left_margin = Inches(1.0)
sec.right_margin = Inches(1.0)
sec.footer_distance = Inches(0.5)
footer=sec.footer.paragraphs[0]
footer.text='Personalization, Privacy, and Trust in AI Robo-Advisory | Page '
add_page_number(footer)

styles=doc.styles
styles['Normal'].font.name='Times New Roman'
styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
styles['Normal'].font.size=Pt(12)
for h in ['Heading 1','Heading 2','Heading 3']:
    styles[h].font.name='Times New Roman'
    styles[h]._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

# ---------------------- title page ----------------------
p=doc.add_paragraph()
p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Personalization, Privacy, and Trust in AI Robo-Advisory:\nExperimental Evidence, Real-World Survey Validation, and Machine-Learning Prediction')
r.bold=True; r.font.size=Pt(18); r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(45,64,89)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Dina A. B. Mahmoud, Fatma A. M. Atia, Maha Moussa, and Doaa G. El-Din')
r.font.size=Pt(12); r.font.name='Times New Roman'; r.bold=True
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Author affiliations and corresponding-author details to be verified and inserted according to the target journal submission system.')
r.font.size=Pt(10); r.italic=True
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run('Manuscript type: empirical multi-study secondary analysis using public and researcher-obtained anonymized survey data')
r.font.size=Pt(11); r.italic=True

doc.add_paragraph()
add_heading(doc,'Abstract', level=1)
abstract=(
"Robo-advisors promise scalable investment guidance, yet adoption depends on more than algorithmic portfolio optimization. "
"Investors must believe that automated advice is useful, trustworthy, understandable, and sufficiently safe with respect to data privacy and security. "
"This manuscript integrates three complementary sources of evidence on AI-enabled robo-advisory design. Study 1 reanalyzes a public 2 x 2 randomized vignette experiment (N = 336) in which respondents evaluated standard versus digital-twin robo-advisors presented through either a dashboard or conversational-AI interface. Digital-twin design substantially increased perceived personalization and privacy concern, while conversational AI increased social presence; both features improved trust, usefulness, and willingness outcomes. Study 2 adds real-world robo-advisor adoption evidence from 428 Indonesian Generation-Z digital-investment-app users. SmartPLS and robust regression analyses show that perceived benefits, trust, performance expectancy, social influence, and awareness channels are positively associated with willingness to use, while risk operates as an adoption inhibitor. An AI/machine-learning extension using five-fold cross-validation predicted high willingness to use robo-advisors with AUC = 0.886 using a random-forest classifier; performance expectancy, perceived benefits, awareness channels, social influence, and trust were the most informative predictors. Study 3 uses the 2024 FINRA National Financial Capability Study Investor Survey (N = 2,861 U.S. investors) to validate the broader digital-investor context. Weighted estimates show that 69.3% of investors trade through a website or mobile app at least sometimes, 77.7% rely on brokerage or advisory tools, and 35.2% worry that Internet investing reduces account security. Machine-learning models predicted digital trading with AUC = 0.877. Together, the evidence supports a balanced design principle: AI personalization and conversational interaction can increase perceived relevance and engagement, but adoption depends on transparent data governance, risk communication, and investor control."
)
add_body_p(doc, abstract)
p=doc.add_paragraph(); p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.DOUBLE
r=p.add_run('Keywords: robo-advisors; financial digital twins; conversational AI; fintech adoption; trust; privacy concern; machine learning; financial literacy; digital investing')
r.font.size=Pt(12); r.italic=True
p=doc.add_paragraph(); p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.DOUBLE
r=p.add_run('JEL classifications: G11; G23; G41; O33; C55')
r.font.size=Pt(12); r.italic=True

doc.add_page_break()

# ---------------------- main manuscript ----------------------
add_heading(doc,'1. Introduction',1)
for text in [
"Robo-advisors have moved from a niche financial-technology product to a central example of algorithmic decision support in consumer finance. By converting a client profile into recommended portfolios, risk allocations, rebalancing rules, and sometimes tax or goal-planning suggestions, robo-advisors can reduce the cost of advice and expand access to investment services. The World Bank (2019) described robo-advice as part of a broader shift toward machine-mediated investment guidance, and empirical research has increasingly examined who adopts these systems, what they trust, and why some investors resist delegating financial judgment to algorithms (Hohenberger et al., 2019; Filiz et al., 2022; Gan et al., 2021).",
"The next generation of robo-advisory platforms is likely to be more personalized and more conversational than earlier questionnaire-driven systems. A financial digital twin can be described as a continuously updated representation of an investor's goals, constraints, income, spending, liabilities, risk capacity, and life events. The digital-twin concept is appealing because it suggests that advice can be more adaptive than static risk-tolerance questionnaires, but it also makes the data requirements of personalized advice more visible (Anshari et al., 2022). At the same time, conversational AI changes the interaction layer. Instead of presenting advice in a static dashboard, the platform can answer questions, explain trade-offs, and simulate a more human-like advisory encounter. The design challenge is therefore two-sided: the same features that can improve relevance and comprehension may also intensify privacy concern, perceived surveillance, and uncertainty about whether the system is acting in the user's interest.",
"The current manuscript addresses this challenge with a multi-study design. The original manuscript focused on a randomized experiment comparing digital-twin and conversational-AI robo-advisor interfaces. That experiment is valuable because it provides causal evidence about design features, but it is based on hypothetical vignettes. Reviewers of a stronger journal submission are likely to ask whether the experimental pattern is supported by real adoption data. The revised manuscript therefore retains the experiment as Study 1 and adds two forms of external validation: a direct real-world robo-advisor survey from Indonesian Generation-Z investment-app users and a nationally relevant U.S. investor survey from the FINRA Foundation's 2024 National Financial Capability Study Investor Survey.",
"The revised contribution is both theoretical and empirical. Theoretically, the paper separates three mechanisms that are often blended together in fintech adoption research: personalization architecture, social-interaction interface design, and investor-level readiness. Empirically, it combines factorial treatment effects, partial least squares structural equation modeling, robust regression, weighted investor-survey estimates, and cross-validated machine-learning prediction. The machine-learning extension is not used as a substitute for theory; rather, it functions as a predictive validity check. If constructs such as benefits, trust, risk, social influence, and privacy/security concern truly organize adoption behavior, they should also carry predictive information in out-of-sample models (Breiman, 2001; Hastie et al., 2009).",
"The findings support a balanced view of AI robo-advisory. In the experiment, digital-twin design increased personalization and privacy concern simultaneously. In the Indonesian adoption survey, perceived benefit was the strongest positive driver of willingness to use, while risk reduced willingness and security/privacy concern strongly predicted perceived risk. In the FINRA investor data, most investors already interact with digital investment channels and advisory tools, but concerns about Internet security and fraud remain substantial. These patterns suggest that the future of robo-advisory adoption depends less on adding AI features indiscriminately and more on designing personalization, explanation, and data governance in ways that convert relevance into trustworthy action."
]: add_body_p(doc,text)

add_heading(doc,'2. Theoretical Background and Hypotheses',1)
add_heading(doc,'2.1 Robo-advisory adoption and algorithmic trust',2)
for text in [
"Technology-adoption theories emphasize perceived usefulness, ease of use, subjective norms, and facilitating conditions as central antecedents of technology use (Davis, 1989; Venkatesh et al., 2003). Financial advice, however, is a high-stakes domain in which trust and perceived risk are especially important. A user may believe that a platform is convenient but still avoid it if investment losses, data misuse, or opaque conflicts of interest seem likely. This is why robo-advisory adoption studies often highlight financial knowledge, trust, usability, performance expectations, and risk perceptions as determinants of intention to use (Hohenberger et al., 2019; Gan et al., 2021; Yi et al., 2023).",
"Algorithm aversion provides an additional explanation for resistance to robo-advisory. Investors may prefer human judgment even when algorithmic recommendations are accurate, because delegation to an algorithm can feel uncomfortable, impersonal, or difficult to contest (Filiz et al., 2022; Greig et al., 2023). In financial planning, this hesitation may be amplified by asymmetry of expertise. Many retail investors cannot independently verify whether a portfolio recommendation is suitable. Therefore, adoption depends not only on the objective quality of the algorithm but also on perceived competence, integrity, fiduciary alignment, and the user's ability to understand the recommendation."
]: add_body_p(doc,text)
add_heading(doc,'2.2 Digital twins, personalization, and privacy calculus',2)
for text in [
"Personalization is a core promise of robo-advisory. Earlier systems typically use questionnaires to infer risk tolerance and investment horizon. A digital-twin approach extends this logic by representing the investor as a dynamic financial profile. The richer profile could incorporate goals, income volatility, household needs, spending patterns, investment history, liabilities, and projected life events. Anshari et al. (2022) describe digital-twin robo-advisory as a potential frontier in personalized financial technology. The expected benefit is improved relevance: advice should feel more tailored, timely, and useful.",
"The same personalization can also trigger privacy calculus. Privacy-calculus models argue that users weigh expected benefits against data-collection risks (Dinev and Hart, 2006; Pavlou, 2003). In robo-advisory, the data are especially sensitive because they may include income, balances, risk tolerance, debt, household information, and future plans. A platform that claims to know the investor deeply may therefore increase both perceived value and perceived exposure. The first two hypotheses retain the original logic of the experiment: H1, digital-twin robo-advisory design increases perceived personalization; and H2, digital-twin design increases privacy concern."
]: add_body_p(doc,text)
add_heading(doc,'2.3 Conversational AI and social presence',2)
for text in [
"Conversational AI changes the way advice is delivered. A dashboard communicates through graphs, tables, and menu structures. A conversational interface uses dialogue, turn-taking, and social cues. The social-presence perspective suggests that such cues can make digital systems feel more responsive and engaging, which may reduce friction in complex decision environments. In financial advice, conversational explanation may be helpful because investors frequently need to interpret risk, fees, diversification, and the consequences of not following advice.",
"The benefit of conversational AI should not be interpreted as a guarantee of better advice. Oehler and Horn (2024) show that large language models and robo-advisors can be compared as sources of financial advice, but the policy challenge is to ensure that users understand what the system can and cannot do. Conversational fluency may create an illusion of competence if the interface sounds human but does not provide auditable reasoning. The third hypothesis is therefore focused on experience rather than objective advice quality: H3, conversational-AI presentation increases perceived social presence relative to a dashboard interface."
]: add_body_p(doc,text)
add_heading(doc,'2.4 From willingness to real adoption and predictive validity',2)
for text in [
"The original experiment measures stated willingness after exposure to a hypothetical vignette. Stated willingness is useful, but reviewers often ask whether the same mechanisms appear in real data. Study 2 addresses this concern using an Indonesian Generation-Z survey dataset that directly measures robo-advisor adoption factors in investment apps. This context is important because young investors often adopt mobile investment applications before they have deep investment experience, and their decisions may be shaped by social influence, digital channels, financial literacy, and privacy/security risk. The Zenodo record states that the dataset was collected to examine drivers and barriers influencing Indonesian Generation-Z users' adoption of robo-advisory features in investment applications (Trinantio and Meyliana, 2026).",
"Study 3 broadens the validation to U.S. investors. The FINRA Investor Survey is not a direct robo-advisor survey, but it is a high-quality investor survey containing digital trading behavior, reliance on advisory tools, Internet-investment security concern, fraud concern, risk tolerance, investment knowledge, and social-media information use. These variables are directly relevant to the adoption environment for AI robo-advisory. The FINRA methodology reports 2,861 adults with non-retirement investments who were primary or shared investment decision-makers in their households (FINRA Foundation, 2025a).",
"The revised hypotheses are therefore organized across studies. H4 states that trust, benefits, performance expectancy, social influence, and awareness channels will be positively associated with robo-advisor willingness, while risk and security/privacy concern will operate as inhibitors. H5 states that machine-learning models using theory-relevant adoption constructs will predict high willingness to use and reported robo-advisor use above chance. H6 states that real investors' digital trading and advisory-tool reliance will coexist with security and fraud concerns, indicating that digital readiness and risk salience are not mutually exclusive."
]: add_body_p(doc,text)

add_fig(doc,'figure1_conceptual_model.png','Figure 1. Integrated conceptual framework linking experimental design features, real-world adoption constructs, investor readiness, and machine-learning validation.',width=6.3)

add_heading(doc,'3. Data and Methods',1)
add_heading(doc,'3.1 Study 1: digital-twin x conversational-AI experiment',2)
for text in [
"Study 1 is a secondary analysis of a public randomized experiment on financial digital twins and conversational interfaces in robo-advisory. The analysis sample includes 336 respondents with investment familiarity. Respondents were assigned to one of four conditions: standard dashboard, standard conversational-AI interface, digital-twin dashboard, or digital-twin conversational-AI interface. The portfolio recommendation was held constant at 60% global equity ETFs, 35% bond ETFs, and 5% cash equivalents, which allows the analysis to isolate the effects of design features rather than differences in advice content (Bonelli, 2026).",
"The main outcomes in Study 1 are perceived personalization, social presence, privacy concern, trust, usefulness, adoption intention, advice compliance, hypothetical allocation willingness, and portfolio acceptance. Digital-twin design and conversational-AI presentation were coded as factorial treatment indicators, and robust ordinary least squares models were estimated with HC3 standard errors. The binary portfolio-acceptance outcome was modeled as a linear probability model for comparability of treatment effects."
]: add_body_p(doc,text)

add_heading(doc,'3.2 Study 2: Indonesian Gen-Z robo-advisor adoption survey',2)
for text in [
"Study 2 uses the Zenodo dataset titled 'Analyzing Factors Enablers and Inhibitors of Robo-Advisor Adoption in Investment Apps by Indonesian Gen-Z' (Trinantio and Meyliana, 2026). The record includes anonymized raw responses, cleaned/processed data, a questionnaire/data dictionary, SmartPLS validity output, and bootstrapping results. The dataset was collected in September to October 2025 and is directly focused on robo-advisor adoption in investment applications.",
"The uploaded raw response file contains 473 rows. After applying the cleaned analysis sheet, the usable construct-level sample contains 428 respondents. Among the 428 respondents who had heard of the feature, 188 reported having used a robo-advisor feature and 240 knew the feature but had not used it. The sample therefore allows both intention-oriented and reported-use analyses. Construct indicators are measured on five-point agreement scales and include social influence (SI), performance expectancy (PE), trust (TR), financial literacy/self-knowledge (FL), information/interface quality (IQ), need for human contact as a barrier (OBU), perceived risk (RISK), perceived benefit (BEN), awareness/adoption channels (AA), willingness to use (WU), and security/privacy concern (SR).",
"Reliability was evaluated with Cronbach's alpha using the raw indicators. SmartPLS measurement output and bootstrapped paths were extracted from the uploaded validity and bootstrap files. To provide a transparent regression complement to the original PLS-SEM approach, standardized robust OLS models were estimated for willingness to use, perceived benefits, and perceived risk. A linear probability model was also estimated for reported robo-advisor use."
]: add_body_p(doc,text)

add_heading(doc,'3.3 Study 3: FINRA NFCS 2024 Investor Survey',2)
for text in [
"Study 3 uses the 2024 FINRA National Financial Capability Study Investor Survey as a real investor validation dataset. The methodology states that the sample consisted of 2,861 adults who completed the State-by-State Survey, indicated that they had investments outside retirement accounts, and were screened as primary or shared investment decision-makers. The survey was self-administered online from July to December 2024 and weighted to approximate the investor population by age and education (FINRA Foundation, 2025a).",
"The FINRA questionnaire contains variables that are highly relevant to digital-advice readiness. It asks how often investors place orders online through a website and through a mobile app, whether they worry that Internet investing makes accounts less secure, whether they worry about investment fraud, how much they rely on brokerage or financial-advisory research tools, how much they rely on financial professionals, whether they rely on popular investments displayed in mobile apps, and whether they use social-media sources for investment information (FINRA Foundation, 2025b). These items do not directly identify robo-advisor use, so Study 3 is interpreted as a digital-investor readiness and risk-salience validation study rather than as a direct robo-advisor-adoption model."
]: add_body_p(doc,text)

add_heading(doc,'3.4 Machine-learning analysis',2)
for text in [
"The machine-learning analysis was added to strengthen the manuscript with predictive evidence. For Study 2, two targets were analyzed: high willingness to use robo-advisors, defined as a willingness score at or above the sample median, and reported robo-advisor use, coded as having used the robo-advisor feature versus knowing it but not having used it. Predictors were the adoption constructs excluding willingness itself. For Study 3, the target was digital trading status, defined as placing orders through a website or mobile app at least sometimes. Predictors included security concern, fraud concern, market fairness belief, reliance on advisory tools, reliance on financial professionals, app popularity cues, social media, risk tolerance, self-rated knowledge, objective investing literacy, age, gender, education, and income.",
"Three supervised classifiers were compared using five-fold stratified cross-validation: regularized logistic regression, random forest, and gradient boosting. Performance was summarized with area under the ROC curve (AUC), accuracy, balanced accuracy, and F1 score. Feature importance was reported for the best-performing model for each target using model-based importance. The purpose was not to claim a production-ready prediction engine, but to test whether theory-relevant constructs carry out-of-sample predictive information."
]: add_body_p(doc,text)

# Table 1
rows_t1=[
    {'Study':'1','Dataset / sample':'Digital-twin robo-advisor experiment; N = 336','Role in manuscript':'Main causal evidence','Main outcomes':'Personalization, social presence, privacy concern, trust, usefulness, adoption/advice willingness'},
    {'Study':'2','Dataset / sample':'Indonesian Gen-Z investment-app survey; N = 428 complete construct records','Role in manuscript':'Direct real-world robo-advisor adoption validation','Main outcomes':'Willingness to use and reported robo-advisor use'},
    {'Study':'3','Dataset / sample':'FINRA NFCS 2024 Investor Survey; N = 2,861','Role in manuscript':'Real investor digital-readiness validation','Main outcomes':'Digital trading, advisory-tool reliance, security concern, fraud concern, investment literacy'}]
add_table(doc, rows_t1, ['Study','Dataset / sample','Role in manuscript','Main outcomes'], 'Table 1. Multi-study data structure and analytic role.', font_size=8)

add_heading(doc,'4. Results',1)
add_heading(doc,'4.1 Study 1: experimental evidence on personalization, privacy, and conversational AI',2)
for text in [
"The experimental results support the personalization-privacy trade-off. Digital-twin design increased perceived personalization by 1.43 scale points (95% CI [1.30, 1.56], p < .001) and increased privacy concern by 0.92 scale points (95% CI [0.71, 1.13], p < .001). Conversational AI increased social presence by 1.46 scale points (95% CI [1.31, 1.61], p < .001). These are large effects on seven-point scales and show that the two design features operated through distinct experiential channels.",
"Downstream outcomes provide evidence that both features can improve adoption-relevant responses. Digital-twin design increased trust, usefulness, and the adoption/advice-compliance index; conversational AI also increased trust, usefulness, and the adoption/advice-compliance index. The interaction was positive for the adoption/advice-compliance index and hypothetical allocation willingness, but not uniformly significant across all outcomes. This pattern is important because it suggests complementarity between what the system knows and how it communicates, while also showing that the combination does not erase privacy concern."
]: add_body_p(doc,text)
add_fig(doc,'figure2_study1_effects.png','Figure 2. Robust factorial effects from Study 1. Digital-twin design mainly shifts personalization and privacy salience; conversational AI mainly shifts social presence.',width=6.3)
rows_t2=[]
for label,b,lo,hi,g in [
    ('Perceived personalization: digital twin',1.43,1.30,1.56,'Primary'),('Privacy concern: digital twin',0.92,0.71,1.13,'Primary'),('Social presence: conversational AI',1.46,1.31,1.61,'Primary'),('Trust: digital twin',0.28,0.14,0.42,'Secondary'),('Trust: conversational AI',0.29,0.15,0.42,'Secondary'),('Usefulness: digital twin',0.67,0.54,0.81,'Secondary'),('Adoption/advice index: interaction',0.22,0.04,0.40,'Secondary')]:
    rows_t2.append({'Outcome / contrast':label,'Effect':f'{b:.2f}','95% CI':f'[{lo:.2f}, {hi:.2f}]','Domain':g})
add_table(doc, rows_t2, ['Outcome / contrast','Effect','95% CI','Domain'], 'Table 2. Selected robust treatment effects from Study 1.', note='Effects are reported in scale points except the original allocation and portfolio-acceptance outcomes, which are omitted here for concision.', font_size=8)

add_heading(doc,'4.2 Study 2: construct reliability and adoption profile',2)
for text in [
"Study 2 provides direct real-world robo-advisor adoption evidence. Among respondents retained in the cleaned construct-level sheet, 43.9% reported having used a robo-advisor feature, while 56.1% knew the feature but had not used it. The sample is therefore not merely a general fintech sample; it contains users and aware non-users of robo-advisory features in digital investment applications.",
"Construct reliability was generally acceptable. Willingness to use had alpha = 0.882, perceived risk alpha = 0.880, security/privacy concern alpha = 0.868, social influence alpha = 0.853, performance expectancy alpha = 0.786, trust alpha = 0.770, and the need-for-human-contact barrier alpha = 0.784. Financial literacy/self-knowledge, information/interface quality, and awareness/adoption channels were around alpha = 0.707-0.710. Perceived benefit was borderline with alpha = 0.685, so benefit-related interpretations should be anchored in both reliability and convergent evidence from the SmartPLS and regression results."
]: add_body_p(doc,text)
add_fig(doc,'figure3_study2_construct_profile.png','Figure 3. Study 2 construct means and reliability. Labels show Cronbach alpha values from the raw indicator matrix.',width=6.3)
rows_rel=[]
for r in rel:
    rows_rel.append({'Construct':r['construct'],'Concept':r['label'],'Items':r['items'],'Mean':fmt_num(r['mean'],2),'SD':fmt_num(r['sd'],2),'Alpha':fmt_num(r['alpha'],3)})
add_table(doc, rows_rel, ['Construct','Concept','Items','Mean','SD','Alpha'], 'Table 3. Study 2 construct descriptive statistics and reliability.', note='Scores are means of observed indicators on five-point response scales.', font_size=7)

add_heading(doc,'4.3 Study 2: structural paths and robust regression',2)
for text in [
"The SmartPLS bootstrapping results show a clear dual-factor pattern. Perceived benefit strongly predicted willingness to use (beta = 0.650, p < .001), while perceived risk negatively predicted willingness to use (beta = -0.096, p = .008). Trust and performance expectancy predicted perceived benefit, and financial literacy/self-knowledge predicted trust. Security/privacy concern was the strongest predictor of perceived risk (beta = 0.605, p < .001), consistent with the privacy-calculus logic that perceived data and security exposure can become a barrier to adoption.",
"The robust OLS complement produced a similar interpretation for willingness. In a standardized model, benefit (b = 0.223, p < .001), trust (b = 0.202, p = .002), performance expectancy (b = 0.223, p < .001), social influence (b = 0.243, p < .001), and awareness/adoption channels (b = 0.116, p = .014) were positively associated with willingness to use. Risk and security/privacy concern were not significant in this fully adjusted OLS model, but risk remained a negative path in the SmartPLS model and security/privacy concern strongly predicted risk. This suggests that security concerns may operate partly through risk perceptions rather than as a direct willingness predictor.",
"The reported-use model should be interpreted cautiously because it is cross-sectional and reported use may affect perceptions rather than only the reverse. Still, it is informative that perceived risk, need for human contact, and security/privacy concern were associated with reported-use status in different directions. Users may have more experience with risks and barriers, while aware non-users may respond differently to privacy and human-advice needs. These associations reinforce the need to distinguish intention, experience, and actual behavior in future field data."
]: add_body_p(doc,text)
add_fig(doc,'figure4_study2_pls_paths.png','Figure 4. Study 2 SmartPLS bootstrapped paths. Positive and negative paths show the dual-factor adoption logic.',width=6.3)
rows_pls=[]
for r in pls:
    if r['path'] in ['BEN -> WU','RISK -> WU','SR -> RISK','TR -> BEN','PE -> BEN','FL -> TR','RISK -> OBU','AA -> RISK']:
        rows_pls.append({'Path':r['path'],'Coef.':fmt_num(r['coef'],3),'95% CI':f'[{fmt_num(r.get("ci_low",0),3)}, {fmt_num(r.get("ci_high",0),3)}]','p':fmt_num(r['p'],3)+pstar(r['p'])})
add_table(doc, rows_pls, ['Path','Coef.','95% CI','p'], 'Table 4. Selected Study 2 SmartPLS bootstrapped path coefficients.', note='Stars: * p < .05, ** p < .01, *** p < .001.', font_size=8)

add_heading(doc,'4.4 Study 2: machine-learning prediction of willingness and reported use',2)
for text in [
"The AI/machine-learning extension strengthened the predictive validity of the adoption constructs. High willingness to use robo-advisors was predicted well across classifiers. The random forest achieved the highest cross-validated performance with AUC = 0.886, accuracy = 0.825, balanced accuracy = 0.820, and F1 = 0.849. Gradient boosting was similar (AUC = 0.881), while regularized logistic regression was slightly lower (AUC = 0.858). The high performance indicates that the adoption constructs are not only statistically associated with willingness in-sample, but also carry meaningful out-of-sample predictive information.",
"Predicting reported robo-advisor use was more difficult, as expected. The best classifier was again the random forest with AUC = 0.684. This lower performance is substantively useful. It implies that cross-sectional perceptions explain stated willingness more strongly than they explain whether a respondent has already used the feature. Actual use is likely shaped by app availability, account type, exposure history, marketing, investment amounts, prior onboarding, and other behavioral frictions that are not fully captured by attitudinal constructs.",
"Feature importance for high willingness emphasized performance expectancy, perceived benefit, awareness/adoption channels, social influence, and trust. This pattern aligns with technology-adoption theory and with the PLS-SEM path model. For actual reported use, perceived risk, need for human contact, awareness channels, and information/interface quality were more informative, suggesting that experienced users and non-users may differ in their exposure to product risks and human-advice preferences."
]: add_body_p(doc,text)
add_fig(doc,'figure5_study2_ml_roc_high_wu.png','Figure 5. Study 2 machine-learning ROC curves for high willingness to use robo-advisors.',width=5.8)
add_fig(doc,'figure6_study2_ml_roc_actual_use.png','Figure 6. Study 2 machine-learning ROC curves for reported robo-advisor use.',width=5.8)
add_fig(doc,'figure7_study2_ml_importance.png','Figure 7. Study 2 model-based feature importance for predicting high willingness to use robo-advisors.',width=6.3)
# Table ML perf for Study2
rows_mlp=[]
for r in ml_perf:
    if r['study']=='Study 2':
        rows_mlp.append({'Target':r['target'],'Model':r['model'],'AUC':fmt_num(r['auc_mean'],3),'Accuracy':fmt_num(r['accuracy_mean'],3),'Balanced accuracy':fmt_num(r['balanced_accuracy_mean'],3),'F1':fmt_num(r['f1_mean'],3)})
add_table(doc, rows_mlp, ['Target','Model','AUC','Accuracy','Balanced accuracy','F1'], 'Table 5. Study 2 cross-validated machine-learning performance.', note='Five-fold stratified cross-validation. High willingness is defined by the sample median of the willingness-to-use construct.', font_size=7)

add_heading(doc,'4.5 Study 3: real investor digital readiness and risk salience',2)
for text in [
"The FINRA Investor Survey validates the broader digital-investor context in which AI robo-advisory operates. Weighted estimates show that 62.9% of investors place orders online through a website at least sometimes, 46.8% place orders through a mobile app at least sometimes, and 69.3% use either a website or mobile app for investment order placement at least sometimes. These estimates indicate that digital investment behavior is already mainstream among investors with non-retirement accounts.",
"Digital readiness coexists with advice reliance. Weighted estimates show that 77.7% of investors rely somewhat or a great deal on investment research and tools provided by a brokerage or financial advisory firm, while 71.6% rely somewhat or a great deal on recommendations from personal financial professionals. This coexistence is important for robo-advisor design. It suggests that digital tools do not simply replace advice relationships; rather, investors combine tools, professionals, and digital information sources.",
"Risk salience is also evident. Approximately 35.2% agree that managing investments through the Internet makes their accounts less secure, and 37.0% agree that they are worried about losing money due to investment fraud. The mean Internet-security concern score is 3.83 on a seven-point scale, while the mean fraud concern score is 3.94. These results strongly support the paper's central claim that digital investment adoption and privacy/security concern can rise together."
]: add_body_p(doc,text)
add_fig(doc,'figure8_study3_nfcs_indicators.png','Figure 8. Study 3 weighted FINRA NFCS indicators showing digital investing, advisory-tool reliance, social influence, and risk concern.',width=6.3)
rows_nf=[]
sel_names=['Either online or app trading sometimes/frequently','Rely on brokerage/advisory research tools somewhat/a great deal','Rely on personal financial professionals somewhat/a great deal','Rely on popular investments shown in trading app somewhat/a great deal','Rely on social media groups/message boards somewhat/a great deal','Agree Internet investing makes accounts less secure (5-7)','Agree worried about investment fraud (5-7)','Objective investing literacy: mean percent correct']
for r in nfcs_ind:
    if r['indicator'] in sel_names:
        rows_nf.append({'Indicator':r['indicator'],'Weighted value':fmt_num(r['value'],1),'Scale':r['scale']})
add_table(doc, rows_nf, ['Indicator','Weighted value','Scale'], 'Table 6. Selected weighted indicators from the FINRA NFCS 2024 Investor Survey.', note='Percentages use WGT1. Mean literacy is the percent of objective investing-knowledge items answered correctly.', font_size=8)

add_heading(doc,'4.6 Study 3: regression and machine-learning prediction of digital trading',2)
for text in [
"The weighted regression model explains a large share of digital trading intensity (R-squared = 0.503). Higher reliance on app popularity cues, higher risk tolerance, higher self-rated investing knowledge, higher fraud concern, reliance on brokerage/advisory tools, reliance on social media groups, and social-media personality use were positively associated with digital trading intensity. Age and reliance on financial professionals were negatively associated, suggesting that digital trading is more common among younger investors and among investors who are less dependent on human financial professionals. Internet-security concern was also negatively associated with digital trading intensity, consistent with the privacy/security barrier emphasized in Studies 1 and 2.",
"The machine-learning results were strong. Digital trading status was predicted with AUC = 0.877 by regularized logistic regression, AUC = 0.864 by random forest, and AUC = 0.865 by gradient boosting. The strongest importance signals included reliance on financial professionals, Internet-security concern, social-media personality use, risk tolerance, reliance on app popularity cues, age, fraud concern, reliance on advisory tools, and social-media group reliance. These predictors are behaviorally interpretable and show that digital trading is structured by both readiness and risk factors.",
"Study 3 should not be overclaimed. It does not directly identify robo-advisor use, and it is U.S.-based rather than emerging-market based. Its value is contextual: it shows that real investors are already using digital trading channels and advisory tools while retaining substantial concerns about online investment security and fraud. That context makes the Study 1 and Study 2 findings more externally meaningful."
]: add_body_p(doc,text)
add_fig(doc,'figure9_study3_wls_coefficients.png','Figure 9. Study 3 weighted regression coefficients for digital trading intensity.',width=6.3)
add_fig(doc,'figure10_study3_ml_importance.png','Figure 10. Study 3 machine-learning feature importance for digital trading status.',width=6.3)
add_fig(doc,'figure11_ml_performance_heatmap.png','Figure 11. Cross-validated AI/machine-learning AUC across Study 2 and Study 3 targets.',width=5.8)
rows_wls=[]
for r in nfcs_wls:
    if r['term'] in ['Internet-security concern','Reliance on financial professionals','Reliance on app popularity cues','Risk tolerance','Age','Fraud concern','Reliance on brokerage/advisory tools','Self-rated investing knowledge','Objective literacy']:
        rows_wls.append({'Predictor':r['term'],'Coef.':fmt_num(r['coef'],3),'95% CI':f'[{fmt_num(r["ci_low"],3)}, {fmt_num(r["ci_high"],3)}]','p':fmt_num(r['p'],3)+pstar(r['p'])})
add_table(doc, rows_wls, ['Predictor','Coef.','95% CI','p'], 'Table 7. Selected Study 3 weighted regression coefficients predicting digital trading intensity.', note='Outcome and continuous predictors are standardized. Regression uses FINRA WGT1 weights and HC3 robust standard errors.', font_size=8)
rows_mlp3=[]
for r in ml_perf:
    if r['study']=='Study 3':
        rows_mlp3.append({'Target':r['target'],'Model':r['model'],'AUC':fmt_num(r['auc_mean'],3),'Accuracy':fmt_num(r['accuracy_mean'],3),'Balanced accuracy':fmt_num(r['balanced_accuracy_mean'],3),'F1':fmt_num(r['f1_mean'],3)})
add_table(doc, rows_mlp3, ['Target','Model','AUC','Accuracy','Balanced accuracy','F1'], 'Table 8. Study 3 cross-validated machine-learning performance.', note='Five-fold stratified cross-validation. Target is online or mobile-app order placement at least sometimes.', font_size=8)

add_heading(doc,'5. General Discussion',1)
for text in [
"The three studies converge on a clear conclusion: AI robo-advisory adoption is best understood as a personalization-benefit-trust process constrained by privacy, security, risk, and human-advice preferences. Study 1 demonstrates this process experimentally. Digital twins make advice feel more personal, but also make data collection more salient. Conversational AI makes the interaction feel more socially present, but it does not automatically solve privacy concern. Study 2 demonstrates the same adoption logic in a real robo-advisor survey: benefits, trust, performance expectancy, social influence, and awareness channels explain willingness, while risk and security/privacy concern form the inhibitory channel. Study 3 shows that real investors already use digital investment channels and advisory tools, but many remain worried about Internet-investment security and fraud.",
"The strongest theoretical contribution is the separation of design features from adoption mechanisms. A digital twin is not merely a more advanced robo-advisor; it is a personalization architecture that changes perceived relevance and data exposure simultaneously. Conversational AI is not merely a more modern interface; it is a social-presence mechanism that may improve engagement and comprehension. Investor readiness is not equivalent to trust; real investors can be digitally active and still concerned about fraud and account security. By integrating these levels, the paper provides a more nuanced model than a single technology-acceptance regression could offer.",
"The machine-learning results add a second contribution. Many adoption studies report significant regression coefficients, but predictive validation is less common. In Study 2, high willingness to use robo-advisors was predicted with strong cross-validated discrimination. In Study 3, digital trading was also predicted with strong discrimination. These models do not establish causality, but they show that the measured constructs and investor characteristics contain practical predictive information. For researchers, this supports the value of combining explanatory models with predictive validation. For platform designers, the feature-importance patterns indicate where interventions may be most useful: benefits, performance expectations, trust, awareness, security reassurance, and responsible handling of social/app cues."
]: add_body_p(doc,text)

add_heading(doc,'5.1 Theoretical implications',2)
for text in [
"First, the findings extend privacy-calculus theory to AI financial advice. Privacy concern is not a generic barrier that simply reduces all adoption outcomes. In the experiment, digital twins increased both personalization and privacy concern. In the Indonesian data, security/privacy concern strongly predicted perceived risk, which then reduced willingness in the PLS-SEM model. In the FINRA data, Internet-security concern was negatively associated with digital trading. Together, these results suggest that privacy and security concerns are not peripheral design issues; they are central adoption mechanisms in AI-mediated finance.",
"Second, the findings clarify the role of trust. Trust predicted perceived benefit and willingness in the Indonesian survey, and trust-related outcomes improved in the experiment. However, trust is not enough by itself. Users also need perceived usefulness, social influence, performance expectations, awareness, and interface quality. This supports a multi-construct adoption model in which trust is necessary but not sufficient.",
"Third, the results caution against treating 'AI' as a single adoption stimulus. Digital twins, conversational AI, brokerage/advisory tools, mobile app cues, social-media personalities, and machine-learning prediction are different mechanisms. Each mechanism can influence perception and behavior in a different way. Future studies should therefore specify which AI feature is being tested and avoid overgeneralizing from one type of financial automation to another."
]: add_body_p(doc,text)

add_heading(doc,'5.2 Practical implications',2)
for text in [
"For robo-advisor designers, the results support four practical recommendations. First, personalization should be explainable. Users should know which data are used, how the data affect recommendations, and what they can edit or delete. Second, conversational interfaces should be designed for comprehension rather than persuasion. The goal should be to clarify risk, diversification, fees, and uncertainty, not merely to make advice feel friendly. Third, platforms should treat privacy and security reassurance as part of the product experience. Account security, data minimization, auditability, and clear consent mechanisms should be visible at the point of advice, not hidden in legal disclosures. Fourth, platforms should monitor social and app-based influence. Popular-investment cues and finfluencer-style recommendation pathways may engage users, but they can also increase susceptibility to herding, fraud, or unsuitable decisions.",
"For regulators and financial educators, the findings suggest that digital-advice literacy should include more than basic investment knowledge. Investors need to understand what automated advice is, what data it uses, how recommendations are generated, what conflicts of interest may exist, and how to evaluate whether the recommendation is suitable. The FINRA evidence shows that advisory tools and digital trading are already widespread among investors, which makes AI financial-advice literacy an immediate rather than future concern."
]: add_body_p(doc,text)

add_heading(doc,'5.3 Limitations and future research',2)
for text in [
"Several limitations should be acknowledged. Study 1 is experimental, but it uses hypothetical vignettes rather than real account-opening or portfolio-allocation behavior. Study 2 is directly about robo-advisor adoption, but it is a cross-sectional survey of Indonesian Generation-Z investment-app users, not a random sample of all investors. Study 3 is a high-quality investor survey, but it does not directly measure robo-advisor use or digital twins. Therefore, the strongest causal evidence comes from Study 1, while Study 2 and Study 3 provide external validation rather than definitive behavioral field evidence.",
"Measurement limitations also matter. The Indonesian dataset has generally acceptable reliabilities, but perceived benefit is borderline and some HTMT values suggest overlap among information quality, performance expectancy, and benefits. These constructs are theoretically related, so some overlap is expected, but future work should refine the measurement model and test whether shorter, more distinct scales improve discriminant validity. The reported-use analysis in Study 2 is also cross-sectional, so it cannot determine whether perceptions cause use or use changes perceptions.",
"Future research should test AI robo-advisory features in field experiments with actual onboarding, disclosure choices, account-linking behavior, and portfolio allocations. A strong next study would randomize transparency, data-control design, and conversational explanation in a working robo-advisor prototype, then measure actual willingness to connect accounts, accept recommendations, and maintain allocations over time. Future work should also test heterogeneity by financial literacy, prior investing experience, gender, risk tolerance, income, trust in institutions, and prior exposure to AI tools."
]: add_body_p(doc,text)

add_heading(doc,'6. Conclusion',1)
for text in [
"This revised manuscript uses a stronger evidence base than the original single-experiment paper. Study 1 shows that digital-twin robo-advisory design increases personalization and privacy concern, while conversational AI increases social presence and supports adoption-related willingness. Study 2 shows that in a real robo-advisor adoption survey, benefits, trust, performance expectations, social influence, awareness, risk, and security/privacy concern structure willingness and reported use. Study 3 shows that real U.S. investors are digitally active and rely on advisory tools, but many remain concerned about Internet-investment security and fraud. The machine-learning extension confirms that these constructs carry practical predictive information.",
"The overall message is not that AI robo-advisory should become more personalized at any cost. The better design principle is responsible personalization: make advice relevant, conversational, and understandable while giving users transparent control over data, risks, and recommendation logic. In financial advice, adoption will depend on trust earned through design, governance, and evidence, not on AI novelty alone."
]: add_body_p(doc,text)

add_heading(doc,'Data Availability Statement',1)
for text in [
"Study 1 uses the public Mendeley Data record 'Digital Twin Robo-Advisor Adoption Dataset V_2' (Bonelli, 2026). Study 2 uses the Zenodo record 'Analyzing Factors Enablers and Inhibitors of Robo-Advisor Adoption in Investment Apps by Indonesian Gen-Z' (Trinantio and Meyliana, 2026). Study 3 uses the FINRA Foundation 2024 National Financial Capability Study Investor Survey. FINRA data are available free of charge under FINRA Foundation terms of use; the raw FINRA dataset should not be redistributed by the authors. The reproducibility package accompanying this manuscript contains analysis scripts and derived tables/figures, but not restricted raw FINRA data."
]: add_body_p(doc,text)
add_heading(doc,'Ethics Statement',1)
add_body_p(doc,"No new human-subjects data were collected for this manuscript. The analyses use anonymized secondary survey data. Study 1 and Study 2 are based on public repositories, and Study 3 uses FINRA NFCS data under the FINRA Foundation's stated terms. Any submission should verify whether the target journal requires institutional review board exemption language for secondary analyses.")
add_heading(doc,'Funding, Conflicts of Interest, and AI-Assisted Tools',1)
add_body_p(doc,"Funding and conflict-of-interest statements should be finalized by the authors before submission. During manuscript preparation, AI-assisted drafting and coding support was used to help organize text, run analyses, and generate figures. The authors are responsible for verifying all analyses, references, interpretations, and final content before submission.")

add_heading(doc,'References',1)
refs=[
"Ajzen, I. (1991). The theory of planned behavior. Organizational Behavior and Human Decision Processes, 50(2), 179-211.",
"Anshari, M., Almunawar, M. N., & Masri, M. (2022). Digital Twin: Financial Technology's Next Frontier of Robo-Advisor. Journal of Risk and Financial Management, 15(4), 163.",
"Benartzi, S., & Thaler, R. H. (2007). Heuristics and biases in retirement savings behavior. Journal of Economic Perspectives, 21(3), 81-104.",
"Bonelli, M. (2026). Digital Twin Robo-Advisor Adoption Dataset V_2 (Version 1) [Data set]. Mendeley Data. https://doi.org/10.17632/39rw5ywj8r.1",
"Breiman, L. (2001). Random forests. Machine Learning, 45, 5-32.",
"Davis, F. D. (1989). Perceived usefulness, perceived ease of use, and user acceptance of information technology. MIS Quarterly, 13(3), 319-340.",
"Dinev, T., & Hart, P. (2006). An extended privacy calculus model for e-commerce transactions. Information Systems Research, 17(1), 61-80.",
"Featherman, M. S., & Pavlou, P. A. (2003). Predicting e-services adoption: A perceived risk facets perspective. International Journal of Human-Computer Studies, 59(4), 451-474.",
"Filiz, I., Judek, J. R., Lorenz, M., & Spiwoks, M. (2022). Algorithm aversion as an obstacle in the establishment of robo advisors. Journal of Risk and Financial Management, 15(8), 353.",
"FINRA Investor Education Foundation. (2025a). 2024 National Financial Capability Study: Investor Survey Methodology. FINRA Investor Education Foundation.",
"FINRA Investor Education Foundation. (2025b). 2024 National Financial Capability Study: Investor Survey Instrument. FINRA Investor Education Foundation.",
"FINRA Investor Education Foundation. (2026). NFCS Data and Downloads. FINRA Foundation.",
"Gan, L. Y., Khan, M. T. I., & Liew, T. W. (2021). Understanding consumer's adoption of financial robo-advisors at the outbreak of the COVID-19 crisis in Malaysia. Financial Planning Review, 4(3), e1127.",
"Greig, F., Ramadorai, T., Rossi, A. G., Utkus, S., & Walther, A. (2023). Algorithm aversion: Theory and evidence from robo-advice. Working paper.",
"Hair, J. F., Hult, G. T. M., Ringle, C. M., Sarstedt, M., Danks, N. P., & Ray, S. (2021). Partial Least Squares Structural Equation Modeling (PLS-SEM) Using R. Springer.",
"Hastie, T., Tibshirani, R., & Friedman, J. (2009). The Elements of Statistical Learning: Data Mining, Inference, and Prediction (2nd ed.). Springer.",
"Hohenberger, C., Lee, C., & Coughlin, J. F. (2019). Acceptance of robo-advisors: Effects of financial experience, affective reactions, and self-enhancement motives. Financial Planning Review, 2(2), e1047.",
"Kim, K. T., Hanna, S. D., & Lee, S. T. (2019). Who are robo-advisor users? Journal of Financial Counseling and Planning, 30(1), 70-82.",
"Oehler, A., & Horn, M. (2024). Does ChatGPT provide better advice than robo-advisors? Finance Research Letters, 60, 104898.",
"Pavlou, P. A. (2003). Consumer acceptance of electronic commerce: Integrating trust and risk with the technology acceptance model. International Journal of Electronic Commerce, 7(3), 101-134.",
"Ringle, C. M., Wende, S., & Becker, J.-M. (2024). SmartPLS 4. SmartPLS GmbH.",
"Trinantio, C., & Meyliana. (2026). Analyzing Factors Enablers and Inhibitors of Robo-Advisor Adoption in Investment Apps by Indonesian Gen-Z (Version 1) [Data set]. Zenodo. https://doi.org/10.5281/zenodo.18449755",
"Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. (2003). User acceptance of information technology: Toward a unified view. MIS Quarterly, 27(3), 425-478.",
"World Bank. (2019). Robo-Advisors: Investing through Machines. Research & Policy Briefs, No. 21. World Bank Group.",
"Yi, T. Z., Rom, N. A. M., Hassan, N. M., & Nordin, N. (2023). The adoption of robo-advisory among millennials in the 21st century. Sustainability, 15(7), 6016."
]
for ref in refs:
    p=doc.add_paragraph()
    p.paragraph_format.left_indent=Inches(0.5)
    p.paragraph_format.first_line_indent=Inches(-0.5)
    p.paragraph_format.line_spacing_rule=WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after=Pt(4)
    r=p.add_run(ref)
    r.font.size=Pt(10); r.font.name='Times New Roman'

add_heading(doc,'Appendix A. Variable Mapping for Real-World Validation Studies',1)
add_body_p(doc,"Appendix A summarizes how the uploaded real-world datasets were used in the revised manuscript. This appendix is included in the main file to make the submission self-contained; journals that require shorter main manuscripts can move it to online supplementary material.")
rows_appendix=[
    {'Dataset':'Study 2 Indonesia','Variable block':'WU','Interpretation':'Willingness to use robo-advisor features','Use in analysis':'Primary willingness outcome and high-willingness ML target'},
    {'Dataset':'Study 2 Indonesia','Variable block':'TR, BEN, PE, SI, AA','Interpretation':'Trust, perceived benefit, performance expectancy, social influence, awareness/adoption channels','Use in analysis':'Enabler constructs in PLS, robust OLS, and ML models'},
    {'Dataset':'Study 2 Indonesia','Variable block':'RISK, SR, OBU','Interpretation':'Perceived risk, security/privacy concern, need for human contact','Use in analysis':'Inhibitor constructs and explanatory predictors'},
    {'Dataset':'Study 2 Indonesia','Variable block':'Awareness/use question','Interpretation':'Used robo-advisor feature versus aware non-user','Use in analysis':'Reported-use classifier and linear probability model'},
    {'Dataset':'Study 3 FINRA','Variable block':'C22_3, C22_4','Interpretation':'Website and mobile-app investment order placement','Use in analysis':'Digital trading readiness outcome'},
    {'Dataset':'Study 3 FINRA','Variable block':'C40, D31','Interpretation':'Internet-investment security concern and investment-fraud concern','Use in analysis':'Risk-salience predictors'},
    {'Dataset':'Study 3 FINRA','Variable block':'F30_1, F30_2, F30_3, F30_9, F40','Interpretation':'Human advice, brokerage/advisory tools, app popularity cues, social media, social-media personality decisions','Use in analysis':'Advice and information-source predictors'},
    {'Dataset':'Study 3 FINRA','Variable block':'G2 and G4-G23 quiz items','Interpretation':'Self-rated and objective investing knowledge','Use in analysis':'Investor sophistication predictors'}]
add_table(doc, rows_appendix, ['Dataset','Variable block','Interpretation','Use in analysis'], 'Appendix Table A1. Variable mapping across real-world validation studies.', font_size=7)

add_heading(doc,'Appendix B. Reproducibility Notes',1)
for text in [
"The reproducibility package contains the Python scripts used to parse the uploaded Excel and CSV files, compute reliability statistics, extract SmartPLS and bootstrap results, run robust regressions, fit cross-validated machine-learning classifiers, and generate figures. The package deliberately does not include raw FINRA data because the FINRA Foundation terms prohibit redistribution. To reproduce Study 3, a researcher should independently download the FINRA NFCS 2024 Investor Survey data from the FINRA Foundation and place the data file in the expected local path.",
"All models are intended for transparent secondary analysis rather than production deployment. Machine-learning results use five-fold stratified cross-validation with a fixed random seed. The figures were generated directly from the analysis outputs and embedded in the manuscript. Authors should verify all final affiliations, funding declarations, conflict-of-interest statements, journal reference style, and any journal-specific reporting checklist before submission."
]: add_body_p(doc,text)

# save
os.makedirs(os.path.dirname(DOCX), exist_ok=True)
doc.save(DOCX)
print(DOCX)
